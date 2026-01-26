"""
설문지 텍스트 추출 및 Excel 생성 모듈
"""

import subprocess
from pathlib import Path
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill


def extract_text_from_file(file_path: str, file_ext: str) -> str:
    """
    다양한 형식의 파일에서 텍스트 추출
    """
    if file_ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    elif file_ext == ".docx":
        # python-docx 사용
        try:
            from docx import Document
            doc = Document(file_path)
            
            text_parts = []
            for para in doc.paragraphs:
                text_parts.append(para.text)
            
            # 테이블 내용도 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except ImportError:
            # fallback: pandoc 사용
            result = subprocess.run(
                ["pandoc", file_path, "-t", "plain"],
                capture_output=True, text=True
            )
            return result.stdout
    
    elif file_ext == ".xlsx":
        import pandas as pd
        
        # 모든 시트 읽기
        xls = pd.ExcelFile(file_path)
        text_parts = []
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
            text_parts.append(f"=== {sheet_name} ===")
            text_parts.append(df.to_string(index=False, header=False))
        
        return "\n".join(text_parts)
    
    elif file_ext == ".pdf":
        # PyPDF2 또는 pdfplumber 사용
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts)
        except ImportError:
            raise ValueError("PDF 처리를 위해 pdfplumber를 설치해주세요: pip install pdfplumber")
    
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {file_ext}")


def generate_excel_from_structure(survey_structure: dict, output_path: str):
    """
    Claude가 분석한 설문지 구조를 웹업용 Excel로 변환
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "설문지"
    
    # 헤더 정의
    headers = [
        "문항번호", "콘솔번호", "프로그래밍 로직", "콘솔 로직", 
        "응답가이드", "검수 로직", "질문유형", "보기유형", "문항"
    ]
    # 보기 컬럼 추가 (최대 35개)
    for i in range(1, 36):
        headers.append(f"보기{i}")
    
    # 1,2행 비움, 3행에 헤더
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
    
    # 데이터 입력 (4행부터)
    questions = survey_structure.get("questions", [])
    
    for row_idx, q in enumerate(questions, start=4):
        ws.cell(row=row_idx, column=1, value=q.get("문항번호", ""))
        ws.cell(row=row_idx, column=2, value=q.get("콘솔번호", ""))
        ws.cell(row=row_idx, column=3, value=q.get("프로그래밍_로직"))
        ws.cell(row=row_idx, column=4, value=q.get("콘솔_로직"))
        ws.cell(row=row_idx, column=5, value=q.get("응답가이드"))
        ws.cell(row=row_idx, column=6, value=q.get("검수_로직"))
        ws.cell(row=row_idx, column=7, value=q.get("질문유형", ""))
        ws.cell(row=row_idx, column=8, value=q.get("보기유형", ""))
        ws.cell(row=row_idx, column=9, value=q.get("문항", ""))
        
        # 보기 입력
        options = q.get("보기", [])
        for opt_idx, opt in enumerate(options):
            if opt_idx < 35:  # 최대 35개
                ws.cell(row=row_idx, column=10 + opt_idx, value=opt)
    
    # 컬럼 너비 조정
    column_widths = {
        'A': 10, 'B': 10, 'C': 20, 'D': 25, 'E': 30, 
        'F': 25, 'G': 10, 'H': 12, 'I': 70
    }
    for col, width in column_widths.items():
        ws.column_dimensions[col].width = width
    
    # 보기 컬럼 너비
    from openpyxl.utils import get_column_letter
    for i in range(10, 45):
        ws.column_dimensions[get_column_letter(i)].width = 30
    
    # 저장
    wb.save(output_path)
    return output_path
